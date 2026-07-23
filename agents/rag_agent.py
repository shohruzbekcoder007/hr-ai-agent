"""
Document RAG agent: PDF / Word / MD / TXT → embed → Chroma → answer + sources.

Independent of Hermes/SQL. Switch embedding model via env + full reindex.
Index path is isolated per (provider, model, dim).
"""

from __future__ import annotations

import hashlib
import json
import logging
import os
import shutil
import threading
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional

logger = logging.getLogger("rag_agent")

_lock = threading.RLock()
_service: Optional["RAGAgentService"] = None

_SUPPORTED_SUFFIXES = {".pdf", ".docx", ".md", ".txt", ".markdown"}


def _env(name: str, default: str = "") -> str:
    raw = os.getenv(name)
    return default if raw is None else raw.strip()


def _env_bool(name: str, default: bool) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() in {"1", "true", "yes", "on"}


def _env_int(name: str, default: int) -> int:
    raw = os.getenv(name)
    if raw is None or not str(raw).strip():
        return default
    try:
        return int(raw)
    except ValueError:
        return default


def is_enabled() -> bool:
    return _env_bool("RAG_ENABLED", True)


def _repo_root() -> Path:
    return Path(__file__).resolve().parent.parent


def _docs_dir() -> Path:
    return Path(_env("RAG_DOCS_DIR") or str(_repo_root() / "data" / "docs"))


def _chroma_root() -> Path:
    # Prefer container path (named volume). Avoid Windows bind-mount SQLite
    # "readonly database" (code 1032) on Docker Desktop.
    default_linux = "/home/appuser/.rag/chroma"
    if Path("/home/appuser").is_dir():
        fallback = default_linux
    else:
        fallback = str(_repo_root() / "data" / "rag" / "chroma")
    return Path(_env("RAG_CHROMA_ROOT") or fallback)


def _prompt_path() -> Path:
    return Path(
        _env("RAG_SYSTEM_PROMPT_PATH")
        or str(_repo_root() / "prompts" / "rag_agent_system.md")
    )


def _now_iso() -> str:
    return datetime.now(timezone.utc).isoformat(timespec="seconds")


def _chunk_id(source: str, page: Any, index: int, text: str) -> str:
    raw = f"{source}|{page}|{index}|{text[:200]}"
    return hashlib.sha256(raw.encode("utf-8", errors="ignore")).hexdigest()[:40]


def _stable_id(*parts: Any) -> str:
    raw = "|".join(str(p) for p in parts)
    return hashlib.sha256(raw.encode("utf-8", errors="ignore")).hexdigest()[:40]


def _chroma_meta(meta: dict[str, Any]) -> dict[str, Any]:
    """Chroma only accepts str/int/float/bool scalars."""
    out: dict[str, Any] = {}
    for k, v in (meta or {}).items():
        if v is None:
            continue
        if isinstance(v, bool):
            out[k] = v
        elif isinstance(v, int):
            out[k] = v
        elif isinstance(v, float):
            out[k] = v
        else:
            s = str(v)
            if len(s) > 1000:
                s = s[:1000]
            out[k] = s
    return out


def load_system_prompt() -> str:
    path = _prompt_path()
    if path.is_file():
        return path.read_text(encoding="utf-8")
    return (
        "Answer only from the provided document context. "
        "If missing, say the answer was not found in the documents."
    )


# ---------------------------------------------------------------------------
# document loaders
# ---------------------------------------------------------------------------


def _load_pdf(path: Path) -> list[dict[str, Any]]:
    from pypdf import PdfReader

    out: list[dict[str, Any]] = []
    reader = PdfReader(str(path))
    for i, page in enumerate(reader.pages):
        try:
            text = page.extract_text() or ""
        except Exception:  # noqa: BLE001
            text = ""
        text = text.strip()
        if not text:
            continue
        out.append(
            {
                "text": text,
                "source": path.name,
                "path": str(path),
                "page": i + 1,
                "file_type": "pdf",
            }
        )
    return out


def _load_docx(path: Path) -> list[dict[str, Any]]:
    from docx import Document

    doc = Document(str(path))
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    # tables
    for table in doc.tables:
        for row in table.rows:
            cells = [((c.text or "").strip()) for c in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)
    text = "\n".join(parts).strip()
    if not text:
        return []
    return [
        {
            "text": text,
            "source": path.name,
            "path": str(path),
            "page": None,
            "file_type": "docx",
        }
    ]


def _load_text(path: Path) -> list[dict[str, Any]]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8", errors="replace")
    text = text.strip()
    if not text:
        return []
    suffix = path.suffix.lower().lstrip(".") or "txt"
    return [
        {
            "text": text,
            "source": path.name,
            "path": str(path),
            "page": None,
            "file_type": suffix,
        }
    ]


def load_file(path: Path) -> list[dict[str, Any]]:
    suf = path.suffix.lower()
    if suf == ".pdf":
        return _load_pdf(path)
    if suf == ".docx":
        return _load_docx(path)
    if suf in {".md", ".txt", ".markdown"}:
        return _load_text(path)
    return []


def discover_files(docs_dir: Path) -> list[Path]:
    if not docs_dir.is_dir():
        return []
    files: list[Path] = []
    for p in sorted(docs_dir.rglob("*")):
        if not p.is_file():
            continue
        if p.name.startswith("."):
            continue
        if p.suffix.lower() in _SUPPORTED_SUFFIXES:
            files.append(p)
    return files


def _split_text(text: str, chunk_size: int, overlap: int) -> list[str]:
    text = (text or "").strip()
    if not text:
        return []
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )
        return [c.strip() for c in splitter.split_text(text) if c.strip()]
    except Exception:  # noqa: BLE001
        # Fallback: simple window
        if len(text) <= chunk_size:
            return [text]
        chunks: list[str] = []
        step = max(1, chunk_size - overlap)
        for i in range(0, len(text), step):
            part = text[i : i + chunk_size].strip()
            if part:
                chunks.append(part)
            if i + chunk_size >= len(text):
                break
        return chunks


class RAGAgentService:
    name = "rag_agent"

    def __init__(self) -> None:
        self._ready = False
        self._last_error: str | None = None
        self._identity: Any = None
        self._embeddings: Any = None
        self._collection: Any = None
        self._client: Any = None
        self._manifest: dict[str, Any] = {}
        self._profiles: list[dict[str, Any]] = []
        self._system_prompt = load_system_prompt()
        self.top_k = _env_int("RAG_TOP_K", 5)
        self.chunk_size = _env_int("RAG_CHUNK_SIZE", 800)
        self.chunk_overlap = _env_int("RAG_CHUNK_OVERLAP", 120)
        self.batch_size = _env_int("RAG_EMBED_BATCH_SIZE", 32)

    @property
    def ready(self) -> bool:
        return self._ready and is_enabled()

    def _index_dir(self) -> Path:
        from agents.embeddings import get_embed_identity

        ident = self._identity or get_embed_identity()
        return _chroma_root() / ident.index_key

    def _manifest_path(self) -> Path:
        return self._index_dir() / "manifest.json"

    def _profiles_path(self) -> Path:
        return self._index_dir() / "document_profiles.json"

    def _read_manifest(self) -> dict[str, Any]:
        path = self._manifest_path()
        if not path.is_file():
            return {}
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            return data if isinstance(data, dict) else {}
        except Exception:  # noqa: BLE001
            return {}

    def _write_manifest(self, data: dict[str, Any]) -> None:
        path = self._manifest_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        tmp = path.with_suffix(".tmp")
        tmp.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        os.replace(tmp, path)
        self._manifest = data

    def _read_profiles(self) -> list[dict[str, Any]]:
        path = self._profiles_path()
        if not path.is_file():
            return []
        try:
            data = json.loads(path.read_text(encoding="utf-8"))
            if isinstance(data, dict):
                profs = data.get("profiles")
                return profs if isinstance(profs, list) else []
            if isinstance(data, list):
                return data
        except Exception:  # noqa: BLE001
            return []
        return []

    def _write_profiles(self, profiles: list[dict[str, Any]]) -> None:
        path = self._profiles_path()
        path.parent.mkdir(parents=True, exist_ok=True)
        payload = {"version": 1, "profiles": profiles, "updated_at": _now_iso()}
        tmp = path.with_suffix(".tmp")
        tmp.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        os.replace(tmp, path)
        self._profiles = profiles

    def _identity_matches_manifest(self) -> bool:
        from agents.embeddings import get_embed_identity

        ident = self._identity or get_embed_identity()
        m = self._manifest or self._read_manifest()
        if not m:
            return False
        return (
            str(m.get("provider")) == ident.provider
            and str(m.get("model")) == ident.model
            and int(m.get("dim") or 0) == int(ident.dim)
        )

    def initialize(self) -> dict[str, Any]:
        with _lock:
            if not is_enabled():
                self._ready = False
                self._last_error = "RAG_ENABLED=false"
                return self.readiness()

            try:
                from agents.embeddings import get_embed_identity, get_embeddings

                self._identity = get_embed_identity()
                self._embeddings = get_embeddings(self._identity)
                self._system_prompt = load_system_prompt()
                self.top_k = _env_int("RAG_TOP_K", 5)
                self.chunk_size = _env_int("RAG_CHUNK_SIZE", 800)
                self.chunk_overlap = _env_int("RAG_CHUNK_OVERLAP", 120)
                self.batch_size = _env_int("RAG_EMBED_BATCH_SIZE", 32)

                index_dir = self._index_dir()
                index_dir.mkdir(parents=True, exist_ok=True)
                # Chroma open can raise Rust PanicException (not subclass of Exception)
                # on corrupt / cross-OS volume indexes — never kill the process.
                self._open_chroma(index_dir)
                self._manifest = self._read_manifest()
                self._profiles = self._read_profiles()

                if self._manifest and not self._identity_matches_manifest():
                    self._last_error = (
                        "Index embed identity mismatch "
                        f"(index={self._manifest.get('provider')}/"
                        f"{self._manifest.get('model')}/d{self._manifest.get('dim')}; "
                        f"env={self._identity.provider}/{self._identity.model}/"
                        f"d{self._identity.dim}). Call POST /v1/docs/reindex."
                    )
                    if _env_bool("RAG_AUTO_REINDEX_ON_MISMATCH", False):
                        logger.warning("%s — auto reindex", self._last_error)
                        return self.reindex(force=True)
                    # Still "ready" to accept reindex; search will refuse
                    self._ready = True
                    return self.readiness()

                self._ready = True
                self._last_error = None
                logger.info(
                    "RAG agent ready provider=%s model=%s dim=%s index=%s chunks=%s",
                    self._identity.provider,
                    self._identity.model,
                    self._identity.dim,
                    index_dir,
                    self._chunk_count(),
                )
                return self.readiness()
            except BaseException as exc:  # noqa: BLE001
                if isinstance(exc, (KeyboardInterrupt, SystemExit)):
                    raise
                self._ready = False
                self._last_error = f"{type(exc).__name__}: {exc}"
                self._client = None
                self._collection = None
                logger.error("RAG initialize failed: %s", self._last_error, exc_info=True)
                return self.readiness()

    def _close_chroma(self) -> None:
        """Release client handles and Chroma process-wide cache before rmtree."""
        self._collection = None
        self._client = None
        try:
            from chromadb.api.shared_system_client import SharedSystemClient

            SharedSystemClient.clear_system_cache()
        except Exception as exc:  # noqa: BLE001
            logger.debug("chroma clear_system_cache: %s", exc)

    def _open_chroma(self, index_dir: Path) -> None:
        """
        Open persistent Chroma. On failure (incl. Rust panic from bad sqlite),
        wipe this identity's directory once and retry — common when a Windows-built
        index is mounted into a Linux container.
        """
        import chromadb
        from chromadb.config import Settings

        index_dir = Path(index_dir).resolve()
        index_dir.mkdir(parents=True, exist_ok=True)
        last_err: BaseException | None = None

        for attempt in range(2):
            try:
                self._close_chroma()
                settings = Settings(
                    anonymized_telemetry=False,
                    allow_reset=True,
                )
                self._client = chromadb.PersistentClient(
                    path=str(index_dir),
                    settings=settings,
                )
                self._collection = self._client.get_or_create_collection(
                    name="docs",
                    metadata={"hnsw:space": "cosine"},
                )
                return
            except BaseException as exc:  # noqa: BLE001 — PanicException is BaseException
                if isinstance(exc, (KeyboardInterrupt, SystemExit)):
                    raise
                last_err = exc
                logger.error(
                    "Chroma open failed attempt=%s path=%s err=%s: %s",
                    attempt + 1,
                    index_dir,
                    type(exc).__name__,
                    exc,
                )
                self._close_chroma()
                if attempt == 0:
                    # Drop corrupt / incompatible store (e.g. host OS vs container)
                    shutil.rmtree(index_dir, ignore_errors=True)
                    index_dir.mkdir(parents=True, exist_ok=True)
                    continue
                break

        raise RuntimeError(
            f"Chroma open failed at {index_dir}: {type(last_err).__name__}: {last_err}. "
            "Delete data/rag/chroma/* and POST /v1/docs/reindex."
        )

    def _chunk_count(self) -> int:
        try:
            if self._collection is None:
                return 0
            return int(self._collection.count())
        except Exception:  # noqa: BLE001
            return 0

    @staticmethod
    def _normalize_text(text: str) -> str:
        """Fold apostrophes so ta'til ≈ tatil for matching."""
        t = (text or "").lower()
        for ch in ("'", "'", "ʻ", "ʼ", "`", "´"):
            t = t.replace(ch, "")
        return t

    @classmethod
    def _tokens(cls, text: str) -> set[str]:
        import re

        t = cls._normalize_text(text)
        return {tok for tok in re.findall(r"[0-9A-Za-zЀ-ӿ]+", t) if len(tok) >= 2}

    def _hybrid_rank(
        self,
        query: str,
        docs: list[Any],
        metas: list[Any],
        dists: list[Any],
    ) -> list[dict[str, Any]]:
        """Blend vector similarity with simple token overlap (FAQ-friendly)."""
        q_norm = self._normalize_text(query)
        q_tok = self._tokens(query)
        # Topic flags after apostrophe-folding
        q_about_leave = any(
            k in q_norm
            for k in ("tatil", "otpusk", "таътил", "отпуск", "leave", "vacation")
        )

        ranked: list[dict[str, Any]] = []
        for i, doc in enumerate(docs):
            meta = metas[i] if i < len(metas) else {}
            dist = dists[i] if i < len(dists) else None
            vec_score = 0.0
            if dist is not None:
                try:
                    vec_score = 1.0 - float(dist)
                except (TypeError, ValueError):
                    vec_score = 0.0
            doc_s = str(doc or "")
            d_norm = self._normalize_text(doc_s)
            d_tok = self._tokens(doc_s)
            if q_tok and d_tok:
                lex = len(q_tok & d_tok) / max(1, len(q_tok))
            else:
                lex = 0.0
            # Substring boosts: ta'til query must still hit "tatil" FAQ lines
            if q_about_leave and "tatil" in d_norm:
                lex = max(lex, 0.55)
            if q_about_leave and any(
                n in d_norm for n in ("15 ish", "21 kalendar", "21 календар", "yiliga 15")
            ):
                lex = max(lex, 0.75)
            ftype = str((meta or {}).get("file_type") or "")
            faq_bonus = (
                0.12 if ftype in {"txt", "md", "markdown"} and lex >= 0.4 else 0.0
            )
            score = round(0.55 * vec_score + 0.45 * lex + faq_bonus, 4)
            ranked.append(
                {
                    "document": doc,
                    "metadata": meta if isinstance(meta, dict) else {},
                    "score": score,
                    "vec_score": round(vec_score, 4),
                    "lex_score": round(lex, 4),
                }
            )
        ranked.sort(key=lambda x: x["score"], reverse=True)
        return ranked

    def _add_batches(
        self,
        all_ids: list[str],
        all_docs: list[str],
        all_metas: list[dict[str, Any]],
    ) -> None:
        total = len(all_docs)
        for start in range(0, total, max(1, self.batch_size)):
            end = min(start + self.batch_size, total)
            batch_docs = all_docs[start:end]
            batch_ids = all_ids[start:end]
            batch_metas = all_metas[start:end]
            vectors = self._embeddings.embed_documents(batch_docs)
            if vectors and len(vectors[0]) != int(self._identity.dim):
                actual = len(vectors[0])
                raise RuntimeError(
                    f"Embedding dimension mismatch: config dim={self._identity.dim} "
                    f"but model returned {actual}. Set RAG_EMBED_DIM={actual} "
                    "and reindex."
                )
            self._collection.add(
                ids=batch_ids,
                documents=batch_docs,
                metadatas=batch_metas,
                embeddings=vectors,
            )

    def reindex(self, force: bool = True) -> dict[str, Any]:
        """Full rebuild of Chroma for the **current** embed identity."""
        del force  # always full rebuild for safety across dim/model
        with _lock:
            if not is_enabled():
                return {
                    "success": False,
                    "error": "RAG_ENABLED=false",
                    "error_code": "disabled",
                }

            t0 = time.time()
            warnings: list[str] = []
            try:
                from agents.embeddings import get_embed_identity, get_embeddings

                self._identity = get_embed_identity()
                self._embeddings = get_embeddings(self._identity)
                index_dir = self._index_dir()
                docs_dir = _docs_dir()
                files = discover_files(docs_dir)
                all_ids: list[str] = []
                all_docs: list[str] = []
                all_metas: list[dict[str, Any]] = []
                profiles: list[dict[str, Any]] = []

                from agents.doc_structure import (
                    build_document_profile,
                    build_toc_text,
                    extract_structure,
                    structure_aware_units,
                )

                for fpath in files:
                    try:
                        page_units = load_file(fpath)
                    except Exception as exc:  # noqa: BLE001
                        warnings.append(f"{fpath.name}: load failed: {exc}")
                        continue
                    if not page_units:
                        warnings.append(f"{fpath.name}: no extractable text")
                        continue

                    # Full text + page map for structure extraction
                    parts: list[str] = []
                    page_map: list[tuple[int, int, int]] = []
                    cursor = 0
                    for unit in page_units:
                        t = unit.get("text") or ""
                        if not t:
                            continue
                        start = cursor
                        parts.append(t)
                        cursor += len(t) + 2  # "\n\n"
                        page_no = unit.get("page")
                        if page_no is None:
                            page_no = -1
                        page_map.append((start, cursor, int(page_no)))
                    full_text = "\n\n".join(parts)
                    if not full_text.strip():
                        warnings.append(f"{fpath.name}: empty after join")
                        continue

                    structure = extract_structure(full_text, filename=fpath.name)
                    page_count = sum(
                        1 for u in page_units if u.get("page") is not None
                    )
                    profile = build_document_profile(
                        source=fpath.name,
                        text=full_text,
                        structure=structure,
                        page_count=page_count or len(page_units),
                    )
                    profiles.append(profile)
                    logger.info(
                        "structure %s quality=%s chapters=%d articles=%d type=%s",
                        fpath.name,
                        structure.structure_quality,
                        len(structure.chapters),
                        len(structure.articles),
                        structure.doc_type,
                    )

                    # --- Document profile chunk (for "nima tartibga soladi") ---
                    prof_text = str(profile.get("profile_text") or profile.get("summary"))
                    all_ids.append(_stable_id(fpath.name, "profile"))
                    all_docs.append(prof_text)
                    all_metas.append(
                        _chroma_meta(
                            {
                                "source": fpath.name,
                                "path": str(fpath),
                                "page": -1,
                                "file_type": fpath.suffix.lower().lstrip(".") or "bin",
                                "chunk_kind": "doc_profile",
                                "doc_id": fpath.name,
                                "article_num": "",
                                "chapter_num": "",
                                "heading_path": "document_profile",
                                "structure_quality": structure.structure_quality,
                                "doc_type": structure.doc_type,
                                "chapter_count": profile.get("chapter_count", 0),
                                "article_count": profile.get("article_count", 0),
                            }
                        )
                    )

                    # --- TOC chunk ---
                    toc_text = build_toc_text(structure, fpath.name)
                    all_ids.append(_stable_id(fpath.name, "toc"))
                    all_docs.append(toc_text)
                    all_metas.append(
                        _chroma_meta(
                            {
                                "source": fpath.name,
                                "path": str(fpath),
                                "page": -1,
                                "file_type": fpath.suffix.lower().lstrip(".") or "bin",
                                "chunk_kind": "toc",
                                "doc_id": fpath.name,
                                "article_num": "",
                                "chapter_num": "",
                                "heading_path": "toc",
                                "structure_quality": structure.structure_quality,
                                "doc_type": structure.doc_type,
                                "chapter_count": profile.get("chapter_count", 0),
                                "article_count": profile.get("article_count", 0),
                            }
                        )
                    )

                    # --- Structure-aware units (modda/bob) or semantic fallback ---
                    structured = structure_aware_units(
                        full_text,
                        structure,
                        source=fpath.name,
                        path=str(fpath),
                        file_type=fpath.suffix.lower().lstrip(".") or "bin",
                        page_map=page_map,
                    )

                    if structured:
                        for idx, unit in enumerate(structured):
                            body = unit["text"]
                            # long articles → sub-split but keep metadata
                            pieces = _split_text(
                                body, self.chunk_size, self.chunk_overlap
                            )
                            if not pieces:
                                pieces = [body[: self.chunk_size]]
                            for j, piece in enumerate(pieces):
                                cid = _stable_id(
                                    fpath.name,
                                    unit.get("chunk_kind"),
                                    unit.get("article_num"),
                                    unit.get("chapter_num"),
                                    j,
                                    piece[:80],
                                )
                                meta = {
                                    "source": fpath.name,
                                    "path": str(fpath),
                                    "page": unit.get("page", -1),
                                    "file_type": unit.get("file_type") or "",
                                    "chunk_kind": unit.get("chunk_kind") or "article",
                                    "doc_id": fpath.name,
                                    "article_num": unit.get("article_num") or "",
                                    "article_title": unit.get("article_title") or "",
                                    "chapter_num": unit.get("chapter_num") or "",
                                    "chapter_title": unit.get("chapter_title") or "",
                                    "heading_path": unit.get("heading_path") or "",
                                    "parent_id": unit.get("parent_id") or "",
                                    "structure_quality": structure.structure_quality,
                                    "doc_type": structure.doc_type,
                                    "chunk_index": j,
                                }
                                all_ids.append(cid)
                                all_docs.append(piece)
                                all_metas.append(_chroma_meta(meta))
                    else:
                        # Unstructured: page/paragraph semantic chunks
                        for unit in page_units:
                            pieces = _split_text(
                                unit["text"], self.chunk_size, self.chunk_overlap
                            )
                            for idx, piece in enumerate(pieces):
                                cid = _chunk_id(
                                    unit["source"], unit.get("page"), idx, piece
                                )
                                meta = {
                                    "source": str(unit["source"]),
                                    "path": str(unit.get("path") or fpath),
                                    "page": unit.get("page")
                                    if unit.get("page") is not None
                                    else -1,
                                    "file_type": str(unit.get("file_type") or ""),
                                    "chunk_kind": "paragraph",
                                    "doc_id": fpath.name,
                                    "article_num": "",
                                    "chapter_num": "",
                                    "heading_path": "",
                                    "structure_quality": structure.structure_quality,
                                    "doc_type": structure.doc_type,
                                    "chunk_index": idx,
                                }
                                all_ids.append(cid)
                                all_docs.append(piece)
                                all_metas.append(_chroma_meta(meta))

                total = len(all_docs)
                # Build into a fresh temp directory, then swap into place. Avoids
                # SQLite "readonly database" when an open Chroma handle still
                # points at a deleted/bind-mounted path (esp. Docker Desktop).
                build_dir = index_dir.parent / f".build_{index_dir.name}_{int(time.time())}"
                if build_dir.exists():
                    shutil.rmtree(build_dir, ignore_errors=True)
                build_dir.mkdir(parents=True, exist_ok=True)

                self._close_chroma()
                self._open_chroma(build_dir)
                self._add_batches(all_ids, all_docs, all_metas)

                self._close_chroma()
                if index_dir.exists():
                    shutil.rmtree(index_dir, ignore_errors=True)
                build_dir.rename(index_dir)
                self._open_chroma(index_dir)

                self._write_profiles(profiles)
                structure_summary = [
                    {
                        "file": p.get("source_file"),
                        "doc_type": p.get("doc_type"),
                        "structure_quality": p.get("structure_quality"),
                        "chapter_count": p.get("chapter_count"),
                        "article_count": p.get("article_count"),
                    }
                    for p in profiles
                ]
                manifest = {
                    "provider": self._identity.provider,
                    "model": self._identity.model,
                    "dim": self._identity.dim,
                    "index_key": self._identity.index_key,
                    "chunk_size": self.chunk_size,
                    "chunk_overlap": self.chunk_overlap,
                    "built_at": _now_iso(),
                    "file_count": len(files),
                    "chunk_count": total,
                    "docs_dir": str(docs_dir),
                    "index_dir": str(index_dir),
                    "files": [p.name for p in files],
                    "document_intelligence": True,
                    "profiles": structure_summary,
                }
                self._write_manifest(manifest)
                self._ready = True
                self._last_error = None
                elapsed = round(time.time() - t0, 2)
                logger.info(
                    "RAG reindex done files=%d chunks=%d dim=%d in %ss path=%s",
                    len(files),
                    total,
                    self._identity.dim,
                    elapsed,
                    index_dir,
                )
                return {
                    "success": True,
                    "file_count": len(files),
                    "chunk_count": total,
                    "elapsed_seconds": elapsed,
                    "warnings": warnings,
                    "identity": self._identity.as_dict(),
                    "index_dir": str(index_dir),
                    "manifest": manifest,
                    "profiles": structure_summary,
                }
            except Exception as exc:  # noqa: BLE001
                self._last_error = str(exc)
                self._ready = False
                self._close_chroma()
                logger.error("RAG reindex failed: %s", exc, exc_info=True)
                return {
                    "success": False,
                    "error": str(exc),
                    "error_code": "reindex_failed",
                    "warnings": warnings,
                }

    def readiness(self) -> dict[str, Any]:
        from agents.embeddings import get_embed_identity

        try:
            ident = self._identity or get_embed_identity()
            ident_d = ident.as_dict()
        except Exception as exc:  # noqa: BLE001
            ident_d = {"error": str(exc)}
            ident = None

        m = self._manifest or self._read_manifest()
        match = False
        try:
            match = self._identity_matches_manifest() if m else False
        except Exception:  # noqa: BLE001
            match = False

        return {
            "ready": bool(self._ready and is_enabled()),
            "enabled": is_enabled(),
            "backend": "rag_chroma",
            "agent": self.name,
            "identity": ident_d,
            "identity_matches_index": match,
            "docs_dir": str(_docs_dir()),
            "chroma_root": str(_chroma_root()),
            "index_dir": str(self._index_dir()) if ident else None,
            "chunk_count": self._chunk_count(),
            "top_k": self.top_k,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "manifest": m or None,
            "profiles_loaded": len(self._profiles or self._read_profiles()),
            "document_intelligence": True,
            "error": self._last_error,
            "supported_types": sorted(_SUPPORTED_SUFFIXES),
        }

    def list_files(self) -> dict[str, Any]:
        docs_dir = _docs_dir()
        files = discover_files(docs_dir)
        m = self._manifest or self._read_manifest()
        indexed = set(m.get("files") or [])
        return {
            "docs_dir": str(docs_dir),
            "files": [
                {
                    "name": p.name,
                    "path": str(p),
                    "suffix": p.suffix.lower(),
                    "size_bytes": p.stat().st_size if p.is_file() else 0,
                    "indexed_in_current_manifest": p.name in indexed,
                }
                for p in files
            ],
            "count": len(files),
        }

    def stats(self) -> dict[str, Any]:
        return self.readiness()

    def chat(self, message: str) -> dict[str, Any]:
        """Answer from retrieved chunks only. Never raises to callers."""
        try:
            message = (message or "").strip()
            if not message:
                return {
                    "success": False,
                    "response": None,
                    "error": "message must not be empty",
                    "error_code": "validation",
                    "sources": [],
                }

            if not is_enabled():
                return {
                    "success": False,
                    "response": None,
                    "error": "RAG disabled (RAG_ENABLED=false)",
                    "error_code": "disabled",
                    "sources": [],
                }

            if not self._ready or self._embeddings is None or self._collection is None:
                self.initialize()
            if not self._ready or self._embeddings is None or self._collection is None:
                return {
                    "success": False,
                    "response": None,
                    "error": self._last_error or "RAG agent not ready",
                    "error_code": "not_ready",
                    "sources": [],
                }

            if self._manifest and not self._identity_matches_manifest():
                return {
                    "success": False,
                    "response": None,
                    "error": self._last_error
                    or "Embed identity mismatch — run POST /v1/docs/reindex",
                    "error_code": "identity_mismatch",
                    "sources": [],
                    "retryable": False,
                }

            count = self._chunk_count()
            if count <= 0:
                return {
                    "success": False,
                    "response": (
                        "Hujjat indeksi bo'sh. PDF/Word fayllarni data/docs ga "
                        "qo'ying va POST /v1/docs/reindex chaqiring."
                    ),
                    "error": "empty_index",
                    "error_code": "empty_index",
                    "sources": [],
                }

            if not self._profiles:
                self._profiles = self._read_profiles()

            from agents.doc_structure import format_counts_answer, route_query

            route = route_query(message)
            mode = "rag_hybrid"
            sources: list[dict[str, Any]] = []
            context_blocks: list[str] = []

            # --- Structured counts (nechta bob/modda) — no embedding needed ---
            if route.wants_counts or "structured_counts" in route.routes:
                answer = format_counts_answer(self._profiles, message)
                for p in self._profiles:
                    sources.append(
                        {
                            "file": p.get("source_file"),
                            "page": None,
                            "score": 1.0,
                            "excerpt": (
                                f"boblar={p.get('chapter_count')}, "
                                f"moddalar={p.get('article_count')}, "
                                f"type={p.get('doc_type')}"
                            ),
                            "file_type": "profile",
                            "chunk_kind": "doc_profile",
                        }
                    )
                return {
                    "success": True,
                    "response": answer,
                    "error": None,
                    "sources": sources,
                    "backend": "rag",
                    "embed_provider": self._identity.provider,
                    "embed_model": self._identity.model,
                    "embed_dim": self._identity.dim,
                    "agents_used": ["rag_agent", "structured_counts"],
                    "mode": "structured_counts",
                    "route": route.routes,
                }

            # --- Metadata filters: article / chapter ---
            where_filter: dict[str, Any] | None = None
            if route.article_num and "article_lookup" in route.routes:
                where_filter = {"article_num": str(route.article_num)}
                mode = "article_lookup"
            elif route.chapter_num and "chapter_lookup" in route.routes:
                where_filter = {"chapter_num": str(route.chapter_num)}
                mode = "chapter_lookup"

            # Hierarchy: which chapter is article N in? answer from metadata
            if "hierarchy" in route.routes and route.article_num:
                try:
                    got = self._collection.get(
                        where={"article_num": str(route.article_num)},
                        include=["documents", "metadatas"],
                    )
                    metas_h = got.get("metadatas") or []
                    docs_h = got.get("documents") or []
                    if metas_h:
                        m0 = metas_h[0] if isinstance(metas_h[0], dict) else {}
                        ch_n = m0.get("chapter_num") or "?"
                        ch_t = m0.get("chapter_title") or ""
                        path = m0.get("heading_path") or ""
                        src = m0.get("source") or ""
                        answer = (
                            f"{route.article_num}-modda "
                            f"**{ch_n}-bob**"
                            + (f" ({ch_t})" if ch_t else "")
                            + f" ichida joylashgan"
                            + (f" [{path}]" if path else "")
                            + (f" — manba: {src}" if src else "")
                            + "."
                        )
                        sources.append(
                            {
                                "file": src,
                                "page": m0.get("page") if m0.get("page") != -1 else None,
                                "score": 1.0,
                                "excerpt": (docs_h[0] or "")[:400] if docs_h else path,
                                "file_type": m0.get("file_type"),
                                "chunk_kind": m0.get("chunk_kind"),
                                "article_num": route.article_num,
                                "chapter_num": ch_n,
                            }
                        )
                        return {
                            "success": True,
                            "response": answer,
                            "error": None,
                            "sources": sources,
                            "backend": "rag",
                            "embed_provider": self._identity.provider,
                            "embed_model": self._identity.model,
                            "embed_dim": self._identity.dim,
                            "agents_used": ["rag_agent", "hierarchy"],
                            "mode": "hierarchy",
                            "route": route.routes,
                        }
                except Exception as exc:  # noqa: BLE001
                    logger.warning("hierarchy lookup failed: %s", exc)

            q_vec = self._embeddings.embed_query(message)
            if len(q_vec) != int(self._identity.dim):
                return {
                    "success": False,
                    "response": None,
                    "error": (
                        f"Query embedding dim {len(q_vec)} != config {self._identity.dim}"
                    ),
                    "error_code": "dim_mismatch",
                    "sources": [],
                }

            fetch_k = max(1, min(max(self.top_k * 4, 12), count))
            query_kwargs: dict[str, Any] = {
                "query_embeddings": [q_vec],
                "n_results": fetch_k,
                "include": ["documents", "metadatas", "distances"],
            }
            if where_filter:
                query_kwargs["where"] = where_filter
                # metadata filter may return fewer hits
                query_kwargs["n_results"] = max(1, min(fetch_k, 20))

            try:
                result = self._collection.query(**query_kwargs)
            except Exception as exc:  # noqa: BLE001
                # where filter may fail if field missing on older index
                logger.warning("chroma query with filter failed (%s); plain query", exc)
                result = self._collection.query(
                    query_embeddings=[q_vec],
                    n_results=fetch_k,
                    include=["documents", "metadatas", "distances"],
                )
                mode = "rag_hybrid"

            docs = (result.get("documents") or [[]])[0]
            metas = (result.get("metadatas") or [[]])[0]
            dists = (result.get("distances") or [[]])[0]
            ranked = self._hybrid_rank(message, docs, metas, dists)

            # Boost profile/toc chunks when route asks for them
            if route.wants_profile or route.wants_toc:
                for item in ranked:
                    kind = str((item.get("metadata") or {}).get("chunk_kind") or "")
                    if route.wants_profile and kind == "doc_profile":
                        item["score"] = float(item.get("score") or 0) + 0.35
                    if route.wants_toc and kind == "toc":
                        item["score"] = float(item.get("score") or 0) + 0.35
                ranked.sort(key=lambda x: float(x.get("score") or 0), reverse=True)

            ranked = ranked[: max(1, self.top_k)]

            for i, item in enumerate(ranked):
                doc = item["document"]
                meta = item["metadata"] or {}
                score = item.get("score")
                page = meta.get("page")
                if page == -1:
                    page = None
                src = {
                    "file": meta.get("source"),
                    "page": page,
                    "score": score,
                    "excerpt": (doc or "")[:400],
                    "file_type": meta.get("file_type"),
                    "chunk_kind": meta.get("chunk_kind"),
                    "article_num": meta.get("article_num"),
                    "chapter_num": meta.get("chapter_num"),
                    "heading_path": meta.get("heading_path"),
                }
                sources.append(src)
                page_s = f" p.{page}" if page is not None else ""
                path_s = meta.get("heading_path") or ""
                path_bit = f" [{path_s}]" if path_s else ""
                context_blocks.append(
                    f"[{i + 1}] source={meta.get('source')}{page_s}{path_bit}\n{doc}"
                )

            if not context_blocks:
                return {
                    "success": True,
                    "response": (
                        "Indekslangan hujjatlardan mos parcha topilmadi. "
                        "Boshqa savol bering yoki hujjatlarni yangilab reindex qiling."
                    ),
                    "sources": [],
                    "backend": "rag",
                    "embed_provider": self._identity.provider,
                    "embed_model": self._identity.model,
                    "embed_dim": self._identity.dim,
                    "agents_used": ["rag_agent"],
                    "mode": mode,
                    "route": route.routes,
                }

            # Inject profile counts into context for LLM when useful
            if self._profiles and (
                route.wants_profile or "doc_profile" in route.routes
            ):
                context_blocks.insert(
                    0,
                    "[profiles]\n" + format_counts_answer(self._profiles, message),
                )

            answer = self._generate_answer(message, "\n\n".join(context_blocks))
            return {
                "success": True,
                "response": answer,
                "error": None,
                "sources": sources,
                "backend": "rag",
                "embed_provider": self._identity.provider,
                "embed_model": self._identity.model,
                "embed_dim": self._identity.dim,
                "agents_used": ["rag_agent", "document_intelligence"],
                "mode": mode,
                "route": route.routes,
            }
        except Exception as exc:  # noqa: BLE001
            logger.error("rag_agent.chat failed: %s", exc, exc_info=True)
            return {
                "success": False,
                "response": None,
                "error": f"RAG agent error: {exc}",
                "error_code": "agent_error",
                "retryable": True,
                "sources": [],
            }

    def _generate_answer(self, question: str, context: str) -> str:
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import HumanMessage, SystemMessage

        model = (
            _env("RAG_LLM_MODEL")
            or _env("LLM_MODEL")
            or _env("OPENAI_MODEL")
            or "gpt-4.1"
        )
        api_key = (
            _env("OPENAI_API_KEY") or _env("LLM_API_KEY") or _env("HERMES_API_KEY")
        )
        base_url = _env("OPENAI_BASE_URL") or None
        kwargs: dict[str, Any] = {
            "model": model,
            "api_key": api_key,
            "temperature": 0,
        }
        if base_url:
            kwargs["base_url"] = base_url
        llm = ChatOpenAI(**kwargs)

        human = (
            f"## Retrieved document context\n\n{context}\n\n"
            f"## User question\n\n{question}\n\n"
            "Answer using only the context above."
        )
        msg = llm.invoke(
            [
                SystemMessage(content=self._system_prompt),
                HumanMessage(content=human),
            ]
        )
        content = getattr(msg, "content", None)
        if isinstance(content, list):
            parts = []
            for b in content:
                if isinstance(b, dict):
                    parts.append(str(b.get("text", b)))
                else:
                    parts.append(str(b))
            return " ".join(parts).strip()
        return str(content if content is not None else msg).strip()


def get_rag_agent() -> RAGAgentService:
    global _service
    with _lock:
        if _service is None:
            _service = RAGAgentService()
            if is_enabled():
                try:
                    _service.initialize()
                except BaseException as exc:  # noqa: BLE001
                    if isinstance(exc, (KeyboardInterrupt, SystemExit)):
                        raise
                    _service._ready = False
                    _service._last_error = f"{type(exc).__name__}: {exc}"
                    logger.error(
                        "get_rag_agent initialize suppressed: %s",
                        _service._last_error,
                    )
        elif is_enabled() and not _service.ready:
            try:
                _service.initialize()
            except BaseException as exc:  # noqa: BLE001
                if isinstance(exc, (KeyboardInterrupt, SystemExit)):
                    raise
                _service._ready = False
                _service._last_error = f"{type(exc).__name__}: {exc}"
                logger.error(
                    "get_rag_agent re-init suppressed: %s", _service._last_error
                )
        return _service
